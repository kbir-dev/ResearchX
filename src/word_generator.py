from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class WordGenerator:
    def __init__(self):
        """Initialize Word document styles"""
        self.document = Document()
        
        # Set up default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set up margins (1 inch on all sides)
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_heading(self, text, level=1):
        """Add a heading with proper formatting"""
        # Add extra space before main headings
        if level == 1:
            self.document.add_paragraph().add_run()
            
        heading = self.document.add_heading('', level)
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18 if level == 1 else 14)
        run.font.bold = (level == 1)  # Only main headings are bold
        
        # Adjust spacing
        heading.space_after = Pt(16 if level == 1 else 12)
        heading.space_before = Pt(24 if level == 1 else 18)
        return heading

    def add_paragraph(self, text, is_methodology_point=False):
        """Add a paragraph with proper formatting"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text.strip())  # Remove extra whitespace
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False
        
        # Adjust indentation for methodology points
        if is_methodology_point:
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent for numbers
        
        # Consistent spacing
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.space_before = Pt(0)  # No space before paragraphs
        paragraph.paragraph_format.line_spacing = 1.5  # 1.5 line spacing
        return paragraph

    def add_reference(self, text):
        """Add a reference with proper formatting"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text.strip())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False
        
        # Hanging indent for references
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.5
        return paragraph

    def markdown_to_word(self, md_file: str, docx_file: str):
        """Convert markdown file to Word document"""
        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            sections = content.split('\n\n')
            current_section = ''
            
            for section in sections:
                if not section.strip():
                    continue
                
                if section.startswith('# '):  # Main heading
                    text = section.replace('# ', '')
                    self.add_heading(text, 1)
                    current_section = text
                    
                elif section.startswith('## '):  # Subheading
                    text = section.replace('## ', '')
                    self.add_heading(text, 2)
                    
                else:  # Normal text
                    clean_text = section.replace('*', '').strip()
                    
                    if 'Methodology' in current_section:
                        lines = clean_text.split('\n')
                        for line in lines:
                            if line.strip():
                                is_point = any(line.strip().startswith(f"{i}.") for i in range(1, 10))
                                self.add_paragraph(line, is_methodology_point=is_point)
                    elif 'Expected Outcomes' in current_section:
                        # Ensure Expected Outcomes text is not bold
                        paragraph = self.document.add_paragraph()
                        run = paragraph.add_run(clean_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = False
                        paragraph.paragraph_format.space_after = Pt(12)
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.5
                    elif 'References' in current_section:
                        self.add_reference(clean_text)
                    else:
                        self.add_paragraph(clean_text)
            
            self.document.save(docx_file)
            return True
            
        except Exception as e:
            print(f"❌ Error converting to Word: {str(e)}")
            raise

def convert_synopsis_to_word():
    """Convert research synopsis from markdown to Word document"""
    print("📄 Converting synopsis to Word document...")
    
    try:
        # Ensure data directory exists
        os.makedirs('data', exist_ok=True)
        
        # Check if markdown file exists
        if not os.path.exists('data/research_synopsis.md'):
            raise FileNotFoundError("Markdown file not found")
            
        word_gen = WordGenerator()
        word_gen.markdown_to_word(
            'data/research_synopsis.md',
            'data/research_synopsis.docx'
        )
        print("✅ Word document generated successfully!")
        
    except FileNotFoundError as e:
        print(f"❌ Error: {str(e)}")
    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}") 