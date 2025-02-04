from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class WordGenerator:
    def __init__(self):
        """Initialize Word document styles"""
        self.document = Document()
        
        # Set up default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set up margins (1 inch on all sides)
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_heading(self, text, level=1):
        """Add a heading with proper formatting"""
        heading = self.document.add_heading('', level)
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18 if level == 1 else 14)
        run.font.bold = True
        heading.space_after = Pt(12)
        return heading

    def add_paragraph(self, text, indent=0):
        """Add a paragraph with proper formatting"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False
        
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(indent * 0.25)
        
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.15
        return paragraph

    def markdown_to_word(self, md_file: str, docx_file: str):
        """Convert markdown file to Word document"""
        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            sections = content.split('\n\n')
            current_section = ''
            
            for section in sections:
                if not section.strip():
                    continue
                
                if section.startswith('# '):  # Main heading
                    text = section.replace('# ', '')
                    self.add_heading(text, 1)
                    current_section = text
                    
                elif section.startswith('### '):  # Subheading
                    text = section.replace('### ', '')
                    self.add_heading(text, 2)
                    
                else:  # Normal text
                    lines = section.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        # Handle Roman numerals
                        if line.startswith('I.') or line.startswith('II.') or line.startswith('III.') or line.startswith('IV.'):
                            self.add_paragraph(line, indent=0)
                        # Handle numbered points
                        elif line.startswith(('1.', '2.', '3.', '4.', '5.')):
                            self.add_paragraph(line, indent=1)
                        else:
                            self.add_paragraph(line)
            
            self.document.save(docx_file)
            return True
            
        except Exception as e:
            print(f"❌ Error converting to Word: {str(e)}")
            raise

def convert_synopsis_to_word():
    """Convert research synopsis from markdown to Word document"""
    print("📄 Converting synopsis to Word document...")
    
    try:
        # Ensure data directory exists
        os.makedirs('data', exist_ok=True)
        
        # Check if markdown file exists
        if not os.path.exists('data/research_synopsis.md'):
            raise FileNotFoundError("Markdown file not found")
            
        word_gen = WordGenerator()
        word_gen.markdown_to_word(
            'data/research_synopsis.md',
            'data/research_synopsis.docx'
        )
        print("✅ Word document generated successfully!")
        
    except FileNotFoundError as e:
        print(f"❌ Error: {str(e)}")
    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}") 