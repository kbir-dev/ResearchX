from typing import List, Dict
import requests
from bs4 import BeautifulSoup
import time
import json
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
import os

class PaperAnalyzer:
    def __init__(self, api_key: str):
        """Initialize with Groq API"""
        try:
            self.client = Groq(
                api_key=api_key,
            )
            print("✅ Connected to Groq API")
        except Exception as e:
            print(f"❌ Error connecting to Groq API: {str(e)}")
            raise
        
    def fetch_paper_content(self, url: str) -> str:
        """Fetch content from paper URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove unnecessary elements
            for element in soup(['script', 'style', 'nav', 'header', 'footer']):
                element.decompose()
                
            return soup.get_text(strip=True)
        except Exception as e:
            print(f"⚠️ Error fetching content from {url}: {str(e)}")
            return ""

    def generate_response(self, prompt: str) -> str:
        """Generate response using Groq API"""
        try:
            # Add timestamp to ensure unique responses
            timestamped_prompt = f"{prompt}\n\nGenerate a unique response. Current time: {time.time()}"
            
            completion = self.client.chat.completions.create(
                model="mixtral-8x7b-32768",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a research assistant skilled in academic writing and analysis. Always provide unique, creative responses while maintaining academic standards."
                    },
                    {
                        "role": "user",
                        "content": timestamped_prompt
                    }
                ],
                temperature=0.8,  # Increased for more variation
                max_tokens=1024,
                top_p=1,
            )
            return completion.choices[0].message.content.strip()
            
        except Exception as e:
            print(f"❌ Error generating response: {str(e)}")
            return "Error generating synopsis. Please try again."

    def analyze_papers(self, papers: List[Dict]) -> List[Dict]:
        """Analyze papers and return structured analysis"""
        print("🔍 Starting paper analysis...")
        
        # Sort and select most relevant papers
        query = papers[0]['Title'].lower()
        papers_sorted = sorted(papers, 
            key=lambda x: sum(word in x['Title'].lower() for word in query.split()))
        relevant_papers = papers_sorted[-3:]
        print(f"📚 Selected {len(relevant_papers)} most relevant papers for analysis")

        # Store papers for references
        self.paper_analyses = relevant_papers
        return relevant_papers

    def create_synopsis(self, papers: List[Dict]) -> dict:
        """Create a research synopsis from papers"""
        print("📝 Generating Research Synopsis...")
        
        try:
            self.paper_analyses = self.analyze_papers(papers)
            research_topic = papers[0]['Title'].lower()
            
            # Generate title
            print("📄 Generating Title...")
            title_prompt = f"Generate a single line, short and concise academic title for research about {research_topic}. Do not include any explanation."
            title = self.generate_response(title_prompt).split('\n')[0]  # Take only the first line
            
            # Generate introduction
            print("📄 Generating Introduction...")
            intro_prompt = (
                f"Write a single comprehensive paragraph (200-250 words) about {research_topic}. "
                "Focus on introducing the problem statement and current challenges in the field. "
                "Keep it formal and academic."
            )
            introduction = self.generate_response(intro_prompt)
            
            # Generate rationale
            print("📄 Generating Rationale...")
            rationale_prompt = (
                f"Write a single focused paragraph (150-200 words) explaining why this research on {research_topic} "
                "is important and required. Focus on the necessity and potential impact. "
                "Keep it concise and well-structured."
            )
            rationale = self.generate_response(rationale_prompt)
            
            # Generate objectives
            print("📄 Generating Objectives...")
            objectives_prompt = (
                f"List exactly 3 one-line, specific objectives for the {research_topic} project. "
                "Format as numbered list (1. 2. 3.). Each objective should start with 'To' and describe "
                "how you will approach the project. Make them clear and achievable."
            )
            objectives = self.generate_response(objectives_prompt)
            
            # Generate literature review
            print("📄 Generating Literature Review...")
            lit_review_prompt = (
                f"Write a literature review (350-400 words) about {research_topic} in 3-4 paragraphs. "
                "Focus on analyzing previous research approaches, their methodologies, and limitations. "
                "Write in flowing paragraphs without citations or references. "
                "Keep it academic and concise, focusing on the approaches rather than specific papers."
            )
            literature_review = self.generate_response(lit_review_prompt)
            
            # Generate feasibility study
            print("📄 Generating Feasibility Study...")
            feasibility_prompt = (
                f"Analyze the feasibility of {research_topic} project in these 4 aspects:\n"
                "I. Technology Feasibility\n"
                "   1. Available technologies and their suitability\n"
                "   2. Technical requirements and implementation\n"
                "II. Financial Feasibility\n"
                "   1. Cost considerations and budget requirements\n"
                "   2. Return on investment analysis\n"
                "III. Time Feasibility\n"
                "   1. Project timeline and milestones\n"
                "   2. Schedule management\n"
                "IV. Resource Feasibility\n"
                "   1. Required resources\n"
                "   2. Resource availability and management\n"
                "Provide a detailed analysis for each aspect without using bullet points. "
                "End with a brief synthesis of the findings."
            )
            feasibility = self.generate_response(feasibility_prompt)
            
            # Generate methodology
            print("📄 Generating Methodology...")
            method_prompt = (
                f"Write a detailed methodology (450-500 words) for {research_topic} in 3-4 paragraphs. "
                "Explain the complete approach including data collection, processing, implementation, "
                "and evaluation methods. Make it technical and specific."
            )
            methodology = self.generate_response(method_prompt)
            
            # Generate facilities required
            print("📄 Generating Facilities Required...")
            facilities_prompt = (
                f"Write a comprehensive list (300-350 words) of technical facilities required for the {research_topic} project using this format:\n"
                "I. Hardware Requirements\n"
                "   1. List specific hardware items\n"
                "   2. Include specifications\n"
                "II. Software Requirements\n"
                "   1. Development environments\n"
                "   2. Frameworks and tools\n"
                "III. Development Tools\n"
                "   1. Testing and deployment tools\n"
                "   2. Version control systems\n"
                "IV. Specialized Equipment\n"
                "   1. List specific equipment\n"
                "   2. Include purpose and specifications\n"
                "Use only Roman numerals for main categories and numbers for subcategories. "
                "Avoid using bullet points, stars, or summary statements."
            )
            facilities = self.generate_response(facilities_prompt)
            
            # Generate expected outcomes
            print("📄 Generating Outcomes...")
            outcomes_prompt = (
                f"Write a detailed section (300-350 words) on expected outcomes after completion of the {research_topic} project. "
                "Include technical achievements, practical applications, and potential impact. "
                "Make it specific and measurable."
            )
            outcomes = self.generate_response(outcomes_prompt)
            
            # Return the synopsis as a structured object
            return {
                "title": title,
                "introduction": introduction,
                "rationale": rationale,
                "objectives": objectives,
                "literature_review": literature_review,
                "feasibility": feasibility,
                "methodology": methodology,
                "facilities": facilities,
                "outcomes": outcomes
            }
            
        except Exception as e:
            print(f"❌ Error creating synopsis: {str(e)}")
            return {
                "title": "Error Generating Synopsis",
                "introduction": "An error occurred while generating the synopsis.",
                "rationale": "",
                "objectives": "",
                "literature_review": "",
                "feasibility": "",
                "methodology": "",
                "facilities": "",
                "outcomes": ""
            }

    def save_synopsis(self, synopsis: Dict[str, str], filename: str = "data/research_synopsis.md"):
        """Save the synopsis to a markdown file"""
        print("💾 Saving Research Synopsis...")
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                # Title only, no explanation
                f.write("# " + synopsis.get('title', '').strip() + "\n\n")
                f.write("---\n\n")
                
                # Introduction section with subsections
                f.write("# Introduction\n\n")
                f.write(synopsis.get('introduction', '') + "\n\n")
                
                # Rationale as a subsection
                f.write("### Rationale\n\n")
                f.write(synopsis.get('rationale', '') + "\n\n")
                
                # Objectives as a subsection
                f.write("### Objectives\n\n")
                f.write(synopsis.get('objectives', '') + "\n\n")
                
                # Main Sections
                main_sections = {
                    'Literature Review': synopsis.get('literature_review', ''),
                    'Feasibility Study': synopsis.get('feasibility', ''),
                    'Methodology/Planning of Project': synopsis.get('methodology', ''),
                    'Facilities Required for Proposed Work': synopsis.get('facilities', ''),
                    'Expected Outcomes': synopsis.get('outcomes', '')
                }
                
                for title, content in main_sections.items():
                    f.write(f"# {title}\n\n")
                    f.write(content + "\n\n")
                
                # References at the end
                f.write("# References\n\n")
                for paper in self.paper_analyses:
                    title = paper['Title']
                    authors = paper.get('Authors', '')
                    year = paper.get('Year', '')
                    url = paper.get('URL', '')
                    
                    ref = f"{authors} ({year}). {title}."
                    if url and url != 'N/A':
                        ref += f" Retrieved from {url}"
                    
                    f.write(f"{ref}\n\n")
            
            print(f"✅ Synopsis saved to {filename}")
            
        except Exception as e:
            print(f"❌ Error saving synopsis: {str(e)}")
            raise

    def save_synopsis_as_docx(self, synopsis: dict, filename: str = "data/research_synopsis.docx") -> str:
        """Save synopsis as a Word document"""
        doc = Document()
        
        # Add title
        doc.add_heading(synopsis['title'], 0)
        
        # Add introduction
        doc.add_heading('Introduction', 1)
        doc.add_paragraph(synopsis['introduction'])
        
        # Add rationale
        doc.add_heading('Rationale', 2)
        doc.add_paragraph(synopsis['rationale'])
        
        # Add objectives
        doc.add_heading('Objectives', 2)
        doc.add_paragraph(synopsis['objectives'])
        
        # Add literature review
        doc.add_heading('Literature Review', 1)
        doc.add_paragraph(synopsis['literature_review'])
        
        # Add feasibility
        doc.add_heading('Feasibility Study', 1)
        doc.add_paragraph(synopsis['feasibility'])
        
        # Add methodology
        doc.add_heading('Methodology', 1)
        doc.add_paragraph(synopsis['methodology'])
        
        # Add facilities
        doc.add_heading('Facilities Required', 1)
        doc.add_paragraph(synopsis['facilities'])
        
        # Add outcomes
        doc.add_heading('Expected Outcomes', 1)
        doc.add_paragraph(synopsis['outcomes'])
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        # Save the document
        doc.save(filename)
        return filename
